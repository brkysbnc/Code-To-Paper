"""
Writer Markdown ciktisini basit Word (.docx) belgesine cevirir.

IEEE iki sutun / sablon stilleri burada yok; kullanici IEEE Word sablonuna yapistirma
veya stilleri sonradan uygulama icin ham yapi elde eder.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import List

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def _style_code_paragraph(doc: Document, text: str) -> None:
    """Kod / diyagram / tablo blogunu Courier ile ekler."""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    para.paragraph_format.left_indent = Pt(12)
    para.paragraph_format.space_after = Pt(6)


def _flush_table_buffer(doc: Document, buf: List[str]) -> None:
    """Ardışık pipe-tablo satirlarini tek blokta yazar."""
    if not buf:
        return
    _style_code_paragraph(doc, "\n".join(buf))
    buf.clear()


def markdown_to_docx_bytes(md: str) -> bytes:
    """
    Basit Markdown satirlarini Word'e donusturur: # basliklar, ``` kod, tablo satirlari.

    Inline **kalin** icin cok sinirli destek: tum paragraf ** ile sariliysa kalin yapilir.
    """
    document = Document()
    in_code = False
    code_lines: List[str] = []
    table_buf: List[str] = []

    def flush_code() -> None:
        if not code_lines:
            return
        _style_code_paragraph(document, "\n".join(code_lines))
        code_lines.clear()

    for raw_line in (md or "").splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buf.append(stripped)
            continue
        _flush_table_buffer(document, table_buf)

        if stripped == "" or stripped == "---":
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=2)
            continue
        if stripped.startswith("#### "):
            document.add_heading(stripped[5:].strip(), level=3)
            continue

        m = re.fullmatch(r"\*\*(.+)\*\*", stripped)
        if m:
            p = document.add_paragraph()
            run = p.add_run(m.group(1))
            run.bold = True
            continue

        p = document.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(11)

    flush_code()
    _flush_table_buffer(document, table_buf)

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()
