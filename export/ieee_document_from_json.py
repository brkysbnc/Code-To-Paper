"""
IEEE ConferenceTemplateIEEE.docx uzerine JSON icerik enjekte eder.

Stil adlari sablonla birebir eslesmeli; yoksa Body Text ile duser.
Govde once tamamen temizlenir (ornek metinler silinir); yalnizca uretilen icerik kalir.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from export.ieee_template_export import (
    append_column_transition_paragraph,
    extract_column_break_sectpr,
    normalize_all_sectpr_cols_space_in_document,
)
from export.ooxml_strict_patch import patch_strict_ooxml_to_opc


def _clear_headers_footers(doc: DocumentObject) -> None:
    """Ust/alt bilgi paragraflarindaki ornek metni bosaltir."""
    for sec in doc.sections:
        for block in (sec.header, sec.footer):
            try:
                for p in block.paragraphs:
                    p.text = ""
            except (ValueError, AttributeError):
                continue


def clear_body_keep_sectpr(doc: DocumentObject) -> None:
    """
    Ana govdedeki tum w:p ve w:tbl (ve diger cocuklari) w:sectPr haricinde siler.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_paragraph_styled(doc: DocumentObject, text: str, style_name: str) -> Any:
    """
    Zorunlu stil ile paragraf ekler; stil yoksa Body Text kullanilir (IEEE kurali: stil her zaman).
    """
    try:
        return doc.add_paragraph(text or "", style=style_name)
    except KeyError:
        return doc.add_paragraph(text or "", style="Body Text")


def add_abstract(doc: DocumentObject, abstract_text: str) -> None:
    """Abstract: 'Abstract—' oneki kalin+italik (em dash), ardindan duz metin."""
    p = doc.add_paragraph(style="Abstract")
    r1 = p.add_run("Abstract\u2014")
    r1.bold = True
    r1.italic = True
    p.add_run(abstract_text or "")


def add_keywords(doc: DocumentObject, keywords_text: str) -> None:
    """Keywords: 'Keywords—' oneki italik, ardindan duz metin."""
    p = doc.add_paragraph(style="Keywords")
    r1 = p.add_run("Keywords\u2014")
    r1.italic = True
    p.add_run(keywords_text or "")


def _add_subtree(doc: DocumentObject, subsections: List[Dict[str, Any]]) -> None:
    """Heading 2 / Heading 3 ve Body Text ile alt agaci yazar."""
    for i, sub in enumerate(subsections):
        if not isinstance(sub, dict):
            continue
        heading_text = f"{chr(65 + i)}. {sub['heading']}"
        add_paragraph_styled(doc, heading_text, "Heading 2")
        body = str(sub.get("body") or "").strip()
        if body:
            p = add_paragraph_styled(doc, body, "Body Text")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for ss in sub.get("subsubsections") or []:
            if not isinstance(ss, dict):
                continue
            add_paragraph_styled(doc, str(ss.get("heading") or ""), "Heading 3")
            sb = str(ss.get("body") or "").strip()
            if sb:
                p3 = add_paragraph_styled(doc, sb, "Body Text")
                p3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def build_ieee_document_bytes(template_path: Path | str, content: Dict[str, Any]) -> bytes:
    """
    Sablonu acar, govdeyi temizler, JSON icerigi sablon stilleriyle yazar.

    content semasi: ieee_json_schema.normalize_ieee_paper_content ciktisi.
    """
    path = Path(template_path)
    raw = path.read_bytes()
    patched = patch_strict_ooxml_to_opc(raw)
    doc = Document(BytesIO(patched))

    col_break_sectpr = extract_column_break_sectpr(doc)
    clear_body_keep_sectpr(doc)
    _clear_headers_footers(doc)

    p_title = add_paragraph_styled(doc, str(content.get("title") or ""), "paper title")
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for author in content.get("authors") or []:
        add_paragraph_styled(doc, str(author), "Author")

    add_abstract(doc, str(content.get("abstract") or ""))
    add_keywords(doc, str(content.get("keywords") or ""))

    if col_break_sectpr is not None:
        append_column_transition_paragraph(doc, col_break_sectpr)

    for section in content.get("sections") or []:
        if not isinstance(section, dict):
            continue
        add_paragraph_styled(doc, str(section.get("heading") or ""), "Heading 1")
        b = str(section.get("body") or "").strip()
        if b:
            pb = add_paragraph_styled(doc, b, "Body Text")
            pb.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _add_subtree(doc, list(section.get("subsections") or []))

    add_paragraph_styled(doc, "Acknowledgment", "Heading 5")
    ack = str(content.get("acknowledgment") or "").strip()
    if ack:
        pa = add_paragraph_styled(doc, ack, "Body Text")
        pa.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    add_paragraph_styled(doc, "References", "Heading 5")
    for ref in content.get("references") or []:
        add_paragraph_styled(doc, str(ref), "references")

    normalize_all_sectpr_cols_space_in_document(doc)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
