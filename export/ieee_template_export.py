"""
IEEE Word stil sablonu: yalnizca stilleri korur; sablon icindeki tum ornek/dummy govde
metnini siler ve RAG/Writer Markdown ciktisini sablon stilleriyle yeniden yazar.

Strict OOXML (purl) sablonlari once OPC URI'lerine cevrilir (ooxml_strict_patch).

Akis ozeti (Markdown -> DOCX):
1) Sablon acilir, author block paragraf/tablo elementleri deepcopy ile govdeden CIKARILIR.
2) Govde tamamen temizlenir (yalnizca final w:sectPr saklanir).
3) Final w:sectPr cols num="2" yapilir (govde 2-sutun layout).
4) write_markdown_with_ieee_styles: paper title (1-col) -> author block (1-col) ->
   1-col continuous break -> 2-col continuous break -> Abstract/Keywords/Bolumler.
"""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from export.ooxml_strict_patch import patch_strict_ooxml_to_opc

logger = logging.getLogger(__name__)


def _style_name(doc: DocumentObject, candidates: List[str]) -> str:
    """Sablon icinde var olan ilk paragraf stil adini dondurur."""
    for name in candidates:
        try:
            _ = doc.styles[name]
            return name
        except KeyError:
            continue
    return "Normal"


def _to_roman_numeral(n: int) -> str:
    """1..3999 arasi tamsayilari Roma rakamina cevirir; IEEE bolum sayilari icin yeterli."""
    if n <= 0:
        return ""
    table: list[tuple[int, str]] = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    parts: list[str] = []
    for value, symbol in table:
        while n >= value:
            parts.append(symbol)
            n -= value
    return "".join(parts)


def _mono_runs(paragraph, text: str) -> None:
    """Kod / diyagram / tablo blogunda Courier kullanir (gomulu stil yoksa)."""
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def extract_column_break_sectpr(doc: DocumentObject):
    """
    Govde paragrafindaki w:pPr / w:sectPr kopyasini dondurur (legacy path icin).

    Yeni akista markdown_to_ieee_template_docx_bytes bu deger yerine kendi 1-col/2-col
    continuous sectPr'larini uretir; geriye donuk uyumluluk icin saklanir.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:p"):
            ppr = child.find(qn("w:pPr"))
            if ppr is not None:
                sp = ppr.find(qn("w:sectPr"))
                if sp is not None:
                    return deepcopy(sp)
    return None


def append_column_transition_paragraph(doc: DocumentObject, sect_pr_el) -> None:
    """Verilen sectPr'i icerek bos paragrafi govde sonuna ekler (legacy column break helper)."""
    break_p = OxmlElement("w:p")
    break_ppr = OxmlElement("w:pPr")
    break_ppr.append(sect_pr_el)
    break_p.append(break_ppr)
    doc.element.body.append(break_p)


def _normalize_w_cols_space_to_twips(cols_el) -> None:
    """w:cols / w:space '36pt' / '18pt' degerini twips tamsayisina cevirir; OOXML twips bekler."""
    if cols_el is None:
        return
    key = qn("w:space")
    val = cols_el.get(key)
    if not val:
        return
    s = str(val).strip()
    if s.endswith("pt"):
        try:
            pt = float(s[:-2].strip())
            twips = max(0, int(round(pt * 20)))
            cols_el.set(key, str(twips))
        except ValueError:
            return


def normalize_all_sectpr_cols_space_in_document(doc: DocumentObject) -> None:
    """Belgedeki tum w:cols ogelerinde w:space pt stringlerini twips'e cevirir."""
    for el in doc.element.body.iter():
        if el.tag == qn("w:cols"):
            _normalize_w_cols_space_to_twips(el)


_AUTHOR_STYLE_KEYWORDS: tuple[str, ...] = ("author", "affiliation", "email")


_MERMAID_KEYWORDS: tuple[str, ...] = (
    "graph TD", "graph LR", "graph RL", "graph TB", "graph BT",
    "flowchart", "sequenceDiagram", "classDiagram", "stateDiagram",
    "erDiagram", "journey", "gantt", "pie",
)


def _line_starts_mermaid(stripped: str) -> bool:
    """Satir Mermaid grammar anahtarlarindan biriyle basliyorsa True doner (fenceless mermaid tespiti)."""
    return any(stripped.startswith(k) for k in _MERMAID_KEYWORDS)


def _w_paragraph_style_id(p_el) -> str:
    """w:p elementinin w:pStyle/w:val degerini kucuk harfle dondurur; yoksa bos string."""
    if p_el is None:
        return ""
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return ""
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        return ""
    return str(pstyle.get(qn("w:val")) or "").strip().lower()


def _is_author_block_paragraph(p_el) -> bool:
    """
    Paragrafin author bloguna ait olup olmadigini belirler.

    Kontrol siralamasi:
    1) Paragraf metninde 'abstract' veya 'keywords' geciyorsa author DEGILDIR (false guard) —
       gercek abstract/keywords paragraflari yanlislikla yazar olarak yakalanmasin diye.
    2) pStyle id'si 'author' / 'affiliation' / 'email' iceriyorsa author bloguna aittir.
    """
    text_l = "".join(t.text for t in p_el.iter(qn("w:t")) if t.text).lower()
    if "abstract" in text_l or "keywords" in text_l:
        return False

    sid = _w_paragraph_style_id(p_el)
    if not sid:
        return False
    return any(key in sid for key in _AUTHOR_STYLE_KEYWORDS)


def _table_contains_author_paragraph(tbl_el) -> bool:
    """
    Tablonun gercek bir author tablosu olup olmadigini, hem hucre metnine hem
    paragraf stiline bakarak belirler. Sablonun yazar bilgi placeholder'lari
    'Given Name', 'Surname', 'organization', 'Affiliation', 'email' anahtarlarini icerir.
    """
    if tbl_el is None:
        return False

    all_text = "".join(t.text for t in tbl_el.iter(qn("w:t")) if t.text).lower()
    keyword_hits = ("given name", "surname", "affiliation", "organization", "email")
    if any(k in all_text for k in keyword_hits):
        return True

    for p_el in tbl_el.iter(qn("w:p")):
        if _is_author_block_paragraph(p_el):
            return True
    return False


def _split_author_paragraph_into_lines(p_el) -> list:
    """
    'line 1: ..., line 2: ..., line 3: ...' formatinda tek paragrafa sikistirilmis yazar
    bilgisini, her 'line N:' parcasi icin ayri w:p elementi haline dondurur.

    Hicbir 'line N:' isareti yoksa (Claude'un soyledigi sessiz veri kaybi kose durumu),
    metnin tamamini tek bir paragraf olarak doner; bu sayede yapilandirilmamis yazar
    paragraflari da kaybolmaz.
    """
    full_text = "".join(t.text for t in p_el.iter(qn("w:t")) if t.text)
    pPr_original = p_el.find(qn("w:pPr"))

    parts = re.split(r"(line\s+\d+\s*:)", full_text, flags=re.IGNORECASE)
    paragraphs: list = []
    current_label = ""
    for part in parts:
        if re.match(r"line\s+\d+\s*:", part, flags=re.IGNORECASE):
            current_label = part
            continue
        line_text = (current_label + part).strip()
        if not line_text:
            current_label = ""
            continue
        new_p = OxmlElement("w:p")
        if pPr_original is not None:
            new_p.append(deepcopy(pPr_original))
        r_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = line_text
        r_el.append(t_el)
        new_p.append(r_el)
        paragraphs.append(new_p)
        current_label = ""

    # CLAUDE FALLBACK: 'line N:' bulunamazsa metnin tamamini tek paragraf olarak koru.
    if not paragraphs and full_text.strip():
        new_p = OxmlElement("w:p")
        if pPr_original is not None:
            new_p.append(deepcopy(pPr_original))
        r_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = full_text.strip()
        r_el.append(t_el)
        new_p.append(r_el)
        paragraphs.append(new_p)

    return paragraphs


def _build_author_ghost_table(authors: list):
    """
    Yazar gruplarini IEEE-stili 2x3 (en fazla 6 yazar) gorunmez tabloya yerlestirir.
    Tum kenarliklar w:val='none' olarak isaretlenir; sayfada yan yana 3 sutun gozukur.
    """
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tblJc = OxmlElement("w:jc")
    tblJc.set(qn("w:val"), "center")
    tblPr.append(tblJc)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "none")
        bd.set(qn("w:sz"), "0")
        bd.set(qn("w:space"), "0")
        bd.set(qn("w:color"), "auto")
        borders.append(bd)
    tblPr.append(borders)
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:val"), "0000")
    tblPr.append(tblLook)
    tbl.append(tblPr)

    tblGrid = OxmlElement("w:tblGrid")
    for _ in range(3):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), "1666")
        tblGrid.append(gc)
    tbl.append(tblGrid)

    def _make_cell(paragraphs=None):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), "1666")
        tcW.set(qn("w:type"), "pct")
        tcPr.append(tcW)
        tc.append(tcPr)
        if paragraphs:
            for p in paragraphs:
                ppr = p.find(qn("w:pPr"))
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    p.insert(0, ppr)
                for existing in list(ppr.findall(qn("w:jc"))):
                    ppr.remove(existing)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                ppr.append(jc)
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:before"), "0")
                spacing.set(qn("w:after"), "0")
                spacing.set(qn("w:line"), "240")
                spacing.set(qn("w:lineRule"), "auto")
                ppr.append(spacing)
                tc.append(p)
        else:
            tc.append(OxmlElement("w:p"))
        return tc

    for row in range(2):
        tr = OxmlElement("w:tr")
        for col in range(3):
            idx = row * 3 + col
            if idx < len(authors):
                tr.append(_make_cell(authors[idx]))
            else:
                tr.append(_make_cell())
        tbl.append(tr)

    return tbl


def extract_author_block_elements(doc: DocumentObject) -> list:
    """
    Author paragraflarini govdeden cikarir; mumkunse bunlari 2x3 gorunmez bir tabloya
    yerlestirip dondurur (IEEE sablon davranisi). Yazar paragraflari 'line N:' parcalarina
    bolunur ve 'line 1:' baslangiclari yeni yazar olarak gruplanir.

    Donus listesi:
      - (varsa) 'Note:' aciklama paragrafi (centered)
      - 1 adet w:tbl (3 sutunlu, kenarliksiz) yazar bilgileri
    Yazar paragrafi bulunamazsa: ESKI DAVRANIS (paragraf listesi) — geriye donuk uyumluluk.
    Sablon zaten yazar tablosu tasiyorsa o tablo deepcopy ile oldugu gibi korunur.
    """
    body = doc.element.body
    raw_author_paragraphs: list = []
    note_p = None

    for child in list(body):
        tag = child.tag
        if tag == qn("w:sectPr"):
            continue

        if tag == qn("w:p") and _is_author_block_paragraph(child):
            copied = deepcopy(child)
            for ppr in copied.iter(qn("w:pPr")):
                for sp in list(ppr.findall(qn("w:sectPr"))):
                    ppr.remove(sp)
            text = "".join(t.text for t in copied.iter(qn("w:t")) if t.text)
            if "Note:" in text:
                note_p = copied
            elif text.strip():
                raw_author_paragraphs.extend(_split_author_paragraph_into_lines(copied))
            body.remove(child)
            continue

        if tag == qn("w:tbl") and _table_contains_author_paragraph(child):
            copied = deepcopy(child)
            for ppr in copied.iter(qn("w:pPr")):
                for sp in list(ppr.findall(qn("w:sectPr"))):
                    ppr.remove(sp)
            body.remove(child)
            logger.info(
                "extract_author_block_elements: TEMPLATE TABLE branch "
                "(template author table detected, copied verbatim)"
            )
            return ([note_p] if note_p is not None else []) + [copied]

    if not raw_author_paragraphs:
        logger.info(
            "extract_author_block_elements: NO AUTHORS branch "
            "(no author paragraphs and no author table found)"
        )
        return [note_p] if note_p is not None else []

    authors: list = []
    current_author: list = []
    for p in raw_author_paragraphs:
        p_text = "".join(t.text for t in p.iter(qn("w:t")) if t.text).strip().lower()
        if p_text.startswith("line 1:") and current_author:
            authors.append(current_author)
            current_author = []
        current_author.append(p)
    if current_author:
        authors.append(current_author)

    authors = authors[:6]
    if not authors:
        logger.info(
            "extract_author_block_elements: FALLBACK branch "
            "(authors found but 'line 1:' grouping failed; returning raw paragraphs)"
        )
        return ([note_p] if note_p is not None else []) + raw_author_paragraphs

    logger.info(
        "extract_author_block_elements: GHOST TABLE branch "
        "(built %d-author 2x3 invisible table)",
        len(authors),
    )
    final_elements: list = []
    if note_p is not None:
        ppr = note_p.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            note_p.insert(0, ppr)
        for existing in list(ppr.findall(qn("w:jc"))):
            ppr.remove(existing)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        ppr.append(jc)
        final_elements.append(note_p)

    final_elements.append(_build_author_ghost_table(authors))
    return final_elements


def make_continuous_sectpr_two_columns(space_twips: int = 360, template_sectpr=None):
    """
    type=continuous + cols num=2 olan yeni bir w:sectPr olusturur (govde 2-sutun gecisi).

    template_sectpr verilirse w:pgSz / w:pgMar deepcopy ile korunur; boylece
    inject edilen continuous bolumler de sablonun sayfa boyutu/marjlariyla cizilir.
    OOXML schema sirasi: type -> pgSz -> pgMar -> cols.
    """
    sectpr = OxmlElement("w:sectPr")

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    sectpr.append(type_el)

    if template_sectpr is not None:
        for tag in ("w:pgSz", "w:pgMar"):
            el = template_sectpr.find(qn(tag))
            if el is not None:
                sectpr.append(deepcopy(el))

    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectpr.append(cols)

    return sectpr


def make_continuous_sectpr_single_column(template_sectpr=None):
    """
    type=continuous + cols num=1 olan yeni bir w:sectPr olusturur (title/author 1-sutun kapanisi).

    template_sectpr verilirse w:pgSz / w:pgMar deepcopy ile korunur.
    OOXML schema sirasi: type -> pgSz -> pgMar -> cols.
    """
    sectpr = OxmlElement("w:sectPr")

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "continuous")
    sectpr.append(type_el)

    if template_sectpr is not None:
        for tag in ("w:pgSz", "w:pgMar"):
            el = template_sectpr.find(qn(tag))
            if el is not None:
                sectpr.append(deepcopy(el))

    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "1")
    cols.set(qn("w:space"), "360")
    sectpr.append(cols)

    return sectpr


def clear_template_body_keep_styles(doc: DocumentObject) -> None:
    """
    Govdedeki tum w:p ve w:tbl ogelerini siler; final w:sectPr saklanir.

    Author block onceden extract_author_block_elements ile cikarildigi icin burada koruma yok;
    boylece dummy paper title / abstract / keywords / govde paragraflari tamamen temizlenir.
    """
    body = doc.element.body
    for child in list(body):
        tag = child.tag
        if tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)


def _clear_header_footer_paragraph_text(doc: DocumentObject) -> None:
    """Ust/alt bilgideki ornek metinleri bosaltir; oge varsa korunur."""
    for sec in doc.sections:
        for block in (sec.header, sec.footer):
            try:
                for p in block.paragraphs:
                    p.text = ""
            except (ValueError, AttributeError):
                continue


def strip_writer_layout_labels(md: str) -> str:
    """Writer ciktisindaki PART 1/2/3 etiket satirlarini kaldirir; TRACEABILITY bloku korunur."""
    lines_out: list[str] = []
    part_re = re.compile(r"^\s*PART\s+[123]\s+[—\-–].*$", re.IGNORECASE)
    for line in (md or "").splitlines():
        if part_re.match(line):
            continue
        lines_out.append(line)
    return "\n".join(lines_out).strip()


def peel_manuscript_title(md: str) -> Tuple[Optional[str], str]:
    """Ilk '# Baslik' satirini makale basligi olarak ayirir; kalan metni dondurur."""
    lines = (md or "").splitlines()
    title: Optional[str] = None
    i0 = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if not s:
            continue
        if s.startswith("# ") and not s.startswith("## "):
            title = s[2:].strip()
            i0 = i + 1
            break
        break
    rest = "\n".join(lines[i0:]).strip()
    return title, rest

def write_markdown_with_ieee_styles(
    doc: DocumentObject,
    md: str,
    *,
    paper_title_override: Optional[str] = None,
    col_break_sectpr=None,
    author_block: Optional[list] = None,
) -> None:
    """
    Temizlenmis sablon uzerine Markdown yazar: paper title, author block, Abstract/Keywords, Bolumler.

    author_block verilirse YENI AKIS:
      paper title -> author block append -> 1-col continuous sectPr break -> 2-col continuous sectPr
      break -> Abstract -> Keywords -> ... (col_break_sectpr yok sayilir).

    author_block None ise LEGACY akis: title append + col_break_sectpr Keywords/Heading onunde inject.

    `## TRACEABILITY` ve TRACEABILITY: tablolari Word'e dokulmez; ## icin Roman numara verilir.
    """
    md = strip_writer_layout_labels(md)
    if "TRACEABILITY:" in md:
        md = md.split("TRACEABILITY:", 1)[0].rstrip()
    md = re.sub(
        r"(?:\r?\n)+---[\s\r\n]*##\s+TRACEABILITY\b(?:\s*\([^)]*\))?[\s\S]*\Z",
        "",
        md,
        flags=re.IGNORECASE,
    ).rstrip()
    md = re.sub(
        r"(?is)(?:\r?\n)+TRACEABILITY:\s*(?:\r?\n)+(?:\|[^\r\n]+\|\r?\n)+[\s\S]*\Z",
        "",
        md,
    ).rstrip()
    title, body_md = peel_manuscript_title(md)
    if paper_title_override:
        title = paper_title_override.strip() or title

    paper_title_style = _style_name(doc, ["paper title", "Paper Title", "paper Title"])
    h1_style = _style_name(doc, ["Heading 1"])
    h2_style = _style_name(doc, ["Heading 2"])
    h3_style = _style_name(doc, ["Heading 3", "Heading 4"])
    body_style = _style_name(doc, ["Body Text", "Normal"])
    abstract_style = _style_name(doc, ["Abstract"])
    keywords_style = _style_name(doc, ["Keywords"])
    ref_heading_style = _style_name(doc, ["Heading 5"])
    ref_entry_style = _style_name(doc, ["references", "References"])

    if title:
        tp = doc.add_paragraph(title, style=paper_title_style)
        tp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if author_block is not None:
        # KRITIK: lxml body.append() final w:sectPr'dan SONRA koyar (yanlis pozisyon).
        # OOXML'de final sectPr body'nin son child'i olmali; tum content ondan onceye girer.
        # python-docx'in doc.add_paragraph davranisi (final sectPr'dan once) ile uyumlu olmasi
        # icin author block ve break paragraflarini da final sectPr'dan once insert ediyoruz.
        body = doc.element.body
        final_sectpr = body.find(qn("w:sectPr"))

        def _insert_before_final(el):
            """Verilen elementi final w:sectPr'dan hemen once yerlestirir; final yoksa sona ekler."""
            if final_sectpr is not None:
                final_sectpr.addprevious(el)
            else:
                body.append(el)

        for el in author_block:
            _insert_before_final(el)

        # 1-col continuous: title+author sectionunu kapat.
        p_close = OxmlElement("w:p")
        p_close_ppr = OxmlElement("w:pPr")
        p_close_ppr.append(make_continuous_sectpr_single_column(template_sectpr=final_sectpr))
        p_close.append(p_close_ppr)
        _insert_before_final(p_close)

        # 2-col continuous: Abstract/Keywords/Bolumlerin yer alacagi 2-sutun sectionu ac.
        p_open = OxmlElement("w:p")
        p_open_ppr = OxmlElement("w:pPr")
        p_open_ppr.append(make_continuous_sectpr_two_columns(360, template_sectpr=final_sectpr))
        p_open.append(p_open_ppr)
        _insert_before_final(p_open)
    else:
        if title:
            doc.add_paragraph("", style=body_style)

    in_code = False
    in_mermaid_fence = False
    in_mermaid_section = False
    in_references = False
    code_lines: List[str] = []
    table_buf: List[str] = []
    # Her Heading 1 (# / ##) sonrasinda ### basliklari A., B., C. ile numaralanir.
    subsection_counter = 0
    # IEEE Heading 1 (## / #) Roman numaralandirmasi — TRACEABILITY haric.
    h1_counter = 0
    col_break_injected = False
    last_h1_text = ""
    first_subheading_seen = False

    def flush_code() -> None:
        if not code_lines:
            return
        p = doc.add_paragraph(style=body_style)
        _mono_runs(p, "\n".join(code_lines))
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        code_lines.clear()

    def flush_table() -> None:
        if not table_buf:
            return
        p = doc.add_paragraph(style=body_style)
        _mono_runs(p, "\n".join(table_buf))
        p.paragraph_format.space_after = Pt(6)
        table_buf.clear()

    for raw_line in body_md.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        # ---- Mermaid filtreleri (FENCED + FENCELESS) ----
        # Fenced ```mermaid ... ``` blogunu komple atla; kod olarak bile yazma.
        if stripped.startswith("```mermaid"):
            in_mermaid_fence = True
            continue
        if in_mermaid_fence:
            if stripped.startswith("```"):
                in_mermaid_fence = False
            continue
        # Fenceless mermaid: 'graph TD', 'flowchart', 'sequenceDiagram' gibi anahtar
        # kelimelerle baslayan satir gorulurse bir sonraki heading'e kadar tum
        # icerigi atla. Code icindeysek tetikleme.
        if (
            not in_code
            and not in_mermaid_section
            and _line_starts_mermaid(stripped)
        ):
            in_mermaid_section = True
            continue
        if in_mermaid_section:
            if stripped.startswith("#"):
                in_mermaid_section = False
                # heading isleminin asagidaki dallarda yapilmasi icin akisa devam et
            else:
                continue

        # ---- References modu: '## References' sonrasi [n] satirlari 'references' stilinde
        if in_references:
            if re.match(r"^\s*\[\d+\]", stripped):
                doc.add_paragraph(stripped, style=ref_entry_style)
                continue
            if stripped == "":
                doc.add_paragraph("", style=body_style)
                continue
            # Yeni heading geldiyse moddan cik; akisa devam et (heading dallari islesin).
            if stripped.startswith("#"):
                in_references = False
            else:
                # References icinde olmayan dolgu satirlari atla; bozulmasin.
                in_references = False

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buf.append(stripped)
            continue
        flush_table()

        if stripped in ("", "---"):
            doc.add_paragraph("", style=body_style)
            continue

        # Abstract / Keywords: IEEE paragraf stilleri (Body Text degil).
        abstract_prefix = "Abstract\u2014"
        if stripped.startswith(abstract_prefix):
            p_abs = doc.add_paragraph(style=abstract_style)
            r_abs = p_abs.add_run(abstract_prefix)
            r_abs.bold = True
            r_abs.italic = True
            p_abs.add_run(stripped[len(abstract_prefix):].lstrip())
            continue

        keywords_prefix = "Keywords\u2014"
        if stripped.startswith(keywords_prefix):
            p_kw = doc.add_paragraph(style=keywords_style)
            r_kw = p_kw.add_run(keywords_prefix)
            r_kw.italic = True
            p_kw.add_run(stripped[len(keywords_prefix):].lstrip())
            # Legacy path icin column break inject; yeni akista (author_block doluyken) atlanir.
            if author_block is None and col_break_sectpr is not None and not col_break_injected:
                append_column_transition_paragraph(doc, col_break_sectpr)
                col_break_injected = True
            continue

        # body_md icindeki tek '#' satirlari daima Heading 1 (paper title peel ile ayrildi).
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_raw = stripped[2:].strip()
            if "TRACEABILITY" in heading_raw.upper():
                continue
            if author_block is None and col_break_sectpr is not None and not col_break_injected:
                append_column_transition_paragraph(doc, col_break_sectpr)
                col_break_injected = True
            subsection_counter = 0
            last_h1_text = re.sub(r'[^\w\s]', '', heading_raw.lower())
            last_h1_text = re.sub(r'\s+', ' ', last_h1_text).strip()
            first_subheading_seen = False
            h1_counter += 1
            heading_text = f"{_to_roman_numeral(h1_counter)}. {heading_raw}"
            p = doc.add_paragraph(heading_text, style=h1_style)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        if stripped.startswith("## "):
            heading_raw = stripped[3:].strip()
            if "TRACEABILITY" in heading_raw.upper():
                continue
            # References baslıği: Roman counter HAREKET ETMEZ; ozel mod (ref_entry_style)
            if heading_raw.lower() == "references":
                in_references = True
                doc.add_paragraph("References", style=ref_heading_style)
                subsection_counter = 0
                continue
            if author_block is None and col_break_sectpr is not None and not col_break_injected:
                append_column_transition_paragraph(doc, col_break_sectpr)
                col_break_injected = True
            subsection_counter = 0
            last_h1_text = re.sub(r'[^\w\s]', '', heading_raw.lower())
            last_h1_text = re.sub(r'\s+', ' ', last_h1_text).strip()
            first_subheading_seen = False
            h1_counter += 1
            heading_text = f"{_to_roman_numeral(h1_counter)}. {heading_raw}"
            p = doc.add_paragraph(heading_text, style=h1_style)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            continue

        if stripped.startswith("### "):
            heading_candidate = stripped[4:].strip()
            if not first_subheading_seen:
                cand_norm = re.sub(r'[^\w\s]', '', heading_candidate.lower())
                cand_norm = re.sub(r'\s+', ' ', cand_norm).strip()
                is_duplicate = False
                if cand_norm == last_h1_text:
                    is_duplicate = True
                elif last_h1_text and (cand_norm in last_h1_text or last_h1_text in cand_norm) and abs(len(cand_norm) - len(last_h1_text)) < 10:
                    is_duplicate = True
                
                if is_duplicate:
                    first_subheading_seen = True
                    continue
            
            first_subheading_seen = True
            if author_block is None and col_break_sectpr is not None and not col_break_injected:
                append_column_transition_paragraph(doc, col_break_sectpr)
                col_break_injected = True
            heading_text = f"{chr(65 + subsection_counter)}. {heading_candidate}"
            subsection_counter += 1
            p = doc.add_paragraph(style=h2_style)
            run = p.add_run(heading_text)
            run.italic = True
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph(stripped[5:].strip(), style=h3_style)
            continue

        m = re.fullmatch(r"\*\*(.+)\*\*", stripped)
        if m:
            p = doc.add_paragraph(style=body_style)
            r = p.add_run(m.group(1))
            r.bold = True
            continue

        # Skip plain-text lines that duplicate the current section heading
        _stripped_norm = re.sub(r'[^\w\s]', '', stripped.lower())
        _stripped_norm = re.sub(r'\s+', ' ', _stripped_norm).strip()
        if last_h1_text and _stripped_norm == last_h1_text:
            continue

        p = doc.add_paragraph(line, style=body_style)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for r in p.runs:
            if r.font.size is None:
                r.font.size = Pt(10)

    flush_code()
    flush_table()


def _force_final_sectpr_two_columns(doc: DocumentObject, space_twips: int = 360) -> None:
    """
    Body altindaki final w:sectPr ogesinde w:cols num=2 zorlar.

    Mevcut w:cols varsa silinir, yenisi w:docGrid'den hemen once eklenir (OOXML schema sirasi).
    docGrid yoksa sectPr sonuna eklenir.
    """
    final_sectpr = doc.element.body.find(qn("w:sectPr"))
    if final_sectpr is None:
        return
    for cols in list(final_sectpr.findall(qn("w:cols"))):
        final_sectpr.remove(cols)
    new_cols = OxmlElement("w:cols")
    new_cols.set(qn("w:num"), "2")
    new_cols.set(qn("w:space"), str(space_twips))
    new_cols.set(qn("w:equalWidth"), "1")
    doc_grid = final_sectpr.find(qn("w:docGrid"))
    if doc_grid is not None:
        final_sectpr.insert(list(final_sectpr).index(doc_grid), new_cols)
    else:
        final_sectpr.append(new_cols)


def markdown_to_ieee_template_docx_bytes(
    template_path: Path | str,
    md: str,
    *,
    paper_title_override: Optional[str] = None,
) -> bytes:
    """
    Sablonu yukler, author block'u extract eder, govdeyi temizler ve Markdown'i 2-sutun layout'la yazar.

    Sira: title (1-col) -> author block (1-col) -> 1-col cont. sectPr -> 2-col cont. sectPr ->
    Abstract/Keywords/Bolumler (2-col). Final sectPr da 2-col yapilir.
    """
    path = Path(template_path)
    raw = path.read_bytes()
    patched = patch_strict_ooxml_to_opc(raw)
    doc = Document(BytesIO(patched))

    author_block = extract_author_block_elements(doc)
    # Legacy callerlar icin yine de hesaplanir; yeni akis bunu kullanmaz.
    col_break_sectpr = extract_column_break_sectpr(doc)

    clear_template_body_keep_styles(doc)
    _force_final_sectpr_two_columns(doc, space_twips=360)
    _clear_header_footer_paragraph_text(doc)
    _strip_numpr_from_heading_styles(doc)

    write_markdown_with_ieee_styles(
        doc,
        md,
        paper_title_override=paper_title_override,
        col_break_sectpr=col_break_sectpr,
        author_block=author_block,
    )

    normalize_all_sectpr_cols_space_in_document(doc)

    tables_in_body = doc.element.body.findall(qn("w:tbl"))
    logger.info(
        "markdown_to_ieee_template_docx_bytes: final tables in body=%d",
        len(tables_in_body),
    )

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def resolve_default_ieee_template() -> Optional[Path]:
    """Oncelik: IEEE_DOCX_TEMPLATE env; sonra docs/templates/ConferenceTemplateIEEE.docx."""
    import os

    try:
        from dotenv import load_dotenv

        load_dotenv()
    except ImportError:
        pass

    env = (os.getenv("IEEE_DOCX_TEMPLATE") or "").strip()
    if env:
        p = Path(env)
        if p.is_file():
            return p
    bundled = Path(__file__).resolve().parent.parent / "docs" / "templates" / "ConferenceTemplateIEEE.docx"
    if bundled.is_file():
        return bundled
    return None


def _strip_numpr_from_heading_styles(doc: DocumentObject) -> None:
    """
    Heading 1 / Heading 2 stillerindeki w:numPr otomatik numaralandirmasini siler.

    Sebep: bizim akista heading metni Markdown tarafinda manuel olarak Roman ('I.', 'II.')
    veya Alfa ('A.', 'B.') prefix ile yaziliyor; sablonun stilinde de native numId aktifse
    Word UI'da 'I. I. Introduction' / 'A. A. Scope' seklinde cift numara gorunur.

    Cozum: SADECE bu iki stilin pPr/numPr'sini defansif olarak temizlemek; body
    paragraflarina ve diger stillere DOKUNULMAZ (markdown'dan gelen numarali listeler
    veya Body Text gibi diger stiller etkilenmez).
    """
    target_style_names = ("Heading 1", "Heading 2", "references")
    for style_name in target_style_names:
        try:
            style_el = doc.styles[style_name]._element
        except KeyError:
            continue
        if style_el is None:
            continue
        ppr = style_el.find(qn("w:pPr"))
        if ppr is None:
            continue
        for numpr in list(ppr.findall(qn("w:numPr"))):
            ppr.remove(numpr)


# Eski isimle cagiran kodlar icin geriye donuk takma ad
append_markdown_to_ieee_document = write_markdown_with_ieee_styles
